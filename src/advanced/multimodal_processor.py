"""
Multi-Modal Document Processing
Handles different document types including PDF, images, tables, and structured data
"""

import fitz  # PyMuPDF for PDF processing
from PIL import Image
import pytesseract
import pandas as pd
import numpy as np
from typing import Dict, List, Any, Optional, Tuple
import logging
import os
import io
import base64
from pathlib import Path

logger = logging.getLogger(__name__)

class MultiModalProcessor:
    def __init__(self):
        """Initialize multi-modal processor"""
        self.supported_types = ['.pdf', '.jpg', '.jpeg', '.png', '.xlsx', '.csv', '.docx', '.txt']
        self.image_formats = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']
        self.table_formats = ['.xlsx', '.csv', '.tsv']
        
        # Configure Tesseract if needed
        try:
            pytesseract.get_tesseract_version()
        except Exception as e:
            logger.warning(f"Tesseract not found: {e}")
    
    def process_document(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Process different document types"""
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        processors = {
            '.pdf': self._process_pdf,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image,
            '.gif': self._process_image,
            '.bmp': self._process_image,
            '.tiff': self._process_image,
            '.xlsx': self._process_excel,
            '.csv': self._process_csv,
            '.tsv': self._process_csv,
            '.docx': self._process_docx,
            '.txt': self._process_text
        }
        
        try:
            result = processors[file_ext](str(file_path), **kwargs)
            result['file_type'] = file_ext
            result['file_name'] = file_path.name
            result['file_size'] = file_path.stat().st_size
            return result
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            return {
                'error': str(e),
                'file_type': file_ext,
                'file_name': file_path.name
            }
    
    def _process_pdf(self, file_path: str, extract_images: bool = True, 
                    extract_tables: bool = True) -> Dict[str, Any]:
        """Extract text, metadata, and images from PDF"""
        try:
            doc = fitz.open(file_path)
            content = {
                'text': '',
                'metadata': doc.metadata,
                'pages': [],
                'images': [],
                'tables': []
            }
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                page_content = {
                    'page_number': page_num + 1,
                    'text': text,
                    'images': [],
                    'tables': []
                }
                
                content['text'] += f"\n--- Page {page_num + 1} ---\n{text}"
                
                # Extract images if requested
                if extract_images:
                    page_images = self._extract_images_from_page(page, page_num)
                    page_content['images'] = page_images
                    content['images'].extend(page_images)
                
                # Extract tables if requested
                if extract_tables:
                    page_tables = self._extract_tables_from_page(page, page_num)
                    page_content['tables'] = page_tables
                    content['tables'].extend(page_tables)
                
                content['pages'].append(page_content)
            
            doc.close()
            return content
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    def _process_image(self, file_path: str, ocr_language: str = 'eng') -> Dict[str, Any]:
        """Extract text from images using OCR"""
        try:
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Perform OCR
            text = pytesseract.image_to_string(image, lang=ocr_language)
            
            # Get image metadata
            metadata = {
                'format': image.format,
                'mode': image.mode,
                'size': image.size,
                'has_transparency': image.mode in ('RGBA', 'LA', 'P')
            }
            
            # Extract additional information
            try:
                # Get confidence scores for OCR
                data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            except:
                avg_confidence = 0
            
            return {
                'text': text.strip(),
                'metadata': metadata,
                'ocr_confidence': avg_confidence,
                'word_count': len(text.split()),
                'line_count': len([line for line in text.split('\n') if line.strip()])
            }
            
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            raise
    
    def _process_excel(self, file_path: str, sheet_names: Optional[List[str]] = None) -> Dict[str, Any]:
        """Process Excel files with multiple sheets"""
        try:
            # Read all sheets or specific ones
            if sheet_names:
                df_dict = pd.read_excel(file_path, sheet_name=sheet_names)
            else:
                df_dict = pd.read_excel(file_path, sheet_name=None)
            
            content = {
                'sheets': {},
                'total_sheets': len(df_dict),
                'sheet_names': list(df_dict.keys())
            }
            
            for sheet_name, sheet_data in df_dict.items():
                sheet_info = {
                    'name': sheet_name,
                    'shape': sheet_data.shape,
                    'columns': list(sheet_data.columns),
                    'data_types': sheet_data.dtypes.to_dict(),
                    'sample_data': sheet_data.head(5).to_dict('records'),
                    'summary_stats': self._get_dataframe_summary(sheet_data)
                }
                
                # Extract text content for search
                text_content = self._dataframe_to_text(sheet_data)
                sheet_info['text_content'] = text_content
                
                content['sheets'][sheet_name] = sheet_info
            
            return content
            
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}")
            raise
    
    def _process_csv(self, file_path: str, delimiter: str = ',', encoding: str = 'utf-8') -> Dict[str, Any]:
        """Process CSV files"""
        try:
            df = pd.read_csv(file_path, delimiter=delimiter, encoding=encoding)
            
            content = {
                'shape': df.shape,
                'columns': list(df.columns),
                'data_types': df.dtypes.to_dict(),
                'sample_data': df.head(10).to_dict('records'),
                'summary_stats': self._get_dataframe_summary(df),
                'text_content': self._dataframe_to_text(df)
            }
            
            return content
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
            raise
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process DOCX files"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            content = {
                'text': '',
                'paragraphs': [],
                'tables': [],
                'metadata': {}
            }
            
            # Extract paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    content['paragraphs'].append({
                        'index': i,
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
                    content['text'] += paragraph.text + '\n'
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                content['tables'].append({
                    'index': i,
                    'data': table_data,
                    'rows': len(table_data),
                    'cols': len(table_data[0]) if table_data else 0
                })
            
            return content
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    def _process_text(self, file_path: str, encoding: str = 'utf-8') -> Dict[str, Any]:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
            
            return {
                'text': text,
                'word_count': len(text.split()),
                'line_count': len(text.split('\n')),
                'char_count': len(text)
            }
            
        except Exception as e:
            logger.error(f"Error processing text {file_path}: {e}")
            raise
    
    def _extract_images_from_page(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Extract images from PDF page"""
        images = []
        try:
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    # Get image data
                    xref = img[0]
                    pix = fitz.Pixmap(page.parent, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        img_data = pix.tobytes("png")
                    else:  # CMYK
                        pix1 = fitz.Pixmap(fitz.csRGB, pix)
                        img_data = pix1.tobytes("png")
                        pix1 = None
                    
                    # Encode as base64 for storage
                    img_base64 = base64.b64encode(img_data).decode()
                    
                    image_info = {
                        'page': page_num + 1,
                        'index': img_index,
                        'xref': xref,
                        'width': img[2],
                        'height': img[3],
                        'colorspace': img[4],
                        'data': img_base64,
                        'size': len(img_data)
                    }
                    images.append(image_info)
                    
                except Exception as e:
                    logger.warning(f"Error extracting image {img_index} from page {page_num}: {e}")
                    continue
            
        except Exception as e:
            logger.warning(f"Error extracting images from page {page_num}: {e}")
        
        return images
    
    def _extract_tables_from_page(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Extract tables from PDF page (simplified implementation)"""
        tables = []
        try:
            # This is a simplified table extraction
            # In production, you might want to use tabula-py or similar
            text = page.get_text()
            
            # Look for table-like patterns (rows with consistent separators)
            lines = text.split('\n')
            potential_table_lines = []
            
            for line in lines:
                # Simple heuristic: lines with multiple tabs or consistent spacing
                if '\t' in line or len(line.split()) > 3:
                    potential_table_lines.append(line)
            
            if len(potential_table_lines) > 2:  # At least 3 rows
                table_data = [line.split('\t') if '\t' in line else line.split() 
                            for line in potential_table_lines]
                
                tables.append({
                    'page': page_num + 1,
                    'data': table_data,
                    'rows': len(table_data),
                    'cols': len(table_data[0]) if table_data else 0
                })
        
        except Exception as e:
            logger.warning(f"Error extracting tables from page {page_num}: {e}")
        
        return tables
    
    def _get_dataframe_summary(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Get summary statistics for a DataFrame"""
        try:
            summary = {
                'numeric_columns': df.select_dtypes(include=[np.number]).columns.tolist(),
                'categorical_columns': df.select_dtypes(include=['object']).columns.tolist(),
                'missing_values': df.isnull().sum().to_dict(),
                'memory_usage': df.memory_usage(deep=True).sum()
            }
            
            # Add descriptive statistics for numeric columns
            if summary['numeric_columns']:
                summary['numeric_stats'] = df[summary['numeric_columns']].describe().to_dict()
            
            return summary
            
        except Exception as e:
            logger.warning(f"Error generating DataFrame summary: {e}")
            return {}
    
    def _dataframe_to_text(self, df: pd.DataFrame) -> str:
        """Convert DataFrame to searchable text"""
        try:
            # Convert to string representation
            text_parts = []
            
            # Add column names
            text_parts.append("Columns: " + ", ".join(df.columns))
            
            # Add sample data
            text_parts.append("Sample data:")
            text_parts.append(df.head(10).to_string())
            
            # Add summary statistics
            if not df.empty:
                text_parts.append("Summary:")
                text_parts.append(df.describe().to_string())
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.warning(f"Error converting DataFrame to text: {e}")
            return str(df)
    
    def extract_text_from_image_bytes(self, image_bytes: bytes, 
                                     image_format: str = 'PNG') -> Dict[str, Any]:
        """Extract text from image bytes (for API usage)"""
        try:
            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(image)
            
            return {
                'text': text.strip(),
                'word_count': len(text.split()),
                'confidence': 0.8  # Placeholder
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from image bytes: {e}")
            return {'text': '', 'error': str(e)}
    
    def process_batch(self, file_paths: List[str], **kwargs) -> Dict[str, Any]:
        """Process multiple files in batch"""
        results = {}
        errors = []
        
        for file_path in file_paths:
            try:
                result = self.process_document(file_path, **kwargs)
                results[file_path] = result
            except Exception as e:
                errors.append({'file': file_path, 'error': str(e)})
                logger.error(f"Error processing {file_path}: {e}")
        
        return {
            'results': results,
            'errors': errors,
            'total_files': len(file_paths),
            'successful': len(results),
            'failed': len(errors)
        }
